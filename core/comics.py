"""eduStudio 漫畫製作核心：file-first manifest、版本、Prompt、QA 與匯出。

本模組刻意不把漫畫資料鎖進資料庫。每個 Project 的漫畫都落在：

    projects/<project_id>/comics/series/<series_id>/series.json
    projects/<project_id>/comics/episodes/<story_id>/<version>/
      manifest.json
      source/*.md
      assets/*
      exports/*
      history/manifest_rXXXX.json

AI 只負責產生草稿；CURRENT 與 release 一律需要明確 QA gate。軟體驗證、mock 圖或
Prompt 完成都不能被升格為正式漫畫完成。
"""
from __future__ import annotations

import base64
import ast
import hashlib
import html
import io
import json
import re
import shutil
import threading
import uuid
import zipfile
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Literal

from pydantic import BaseModel, Field, field_validator

from core import config
from core.project import safe_id


EpisodeState = Literal[
    "BACKLOG", "BRIEF", "STORYBOARD", "IMAGE", "LAYOUT", "QA", "CURRENT", "HOLD"
]
QAResult = Literal["PASS", "FAIL", "HOLD", "UNVERIFIED"]
AssetStatus = Literal["DRAFT", "FINAL", "REJECTED"]

STATE_ORDER: tuple[str, ...] = (
    "BACKLOG", "BRIEF", "STORYBOARD", "IMAGE", "LAYOUT", "QA", "CURRENT"
)
REQUIRED_QA_GATES: tuple[str, ...] = (
    "anatomy",
    "technical",
    "text",
    "safety",
    "page_render",
    "human_approval",
)


def utc_now() -> str:
    return datetime.now(timezone.utc).isoformat()


def sha256_file(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as fh:
        for chunk in iter(lambda: fh.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


class Character(BaseModel):
    character_id: str
    name: str
    role: str = ""
    visual_lock: str = ""
    voice: str = ""
    anchor_assets: list[str] = Field(default_factory=list)


class GlossaryTerm(BaseModel):
    term: str
    definition: str
    aliases: list[str] = Field(default_factory=list)


class Series(BaseModel):
    project_id: str
    series_id: str
    title: str
    description: str = ""
    status: Literal["ACTIVE", "PAUSED", "ARCHIVED"] = "ACTIVE"
    visual_bible: str = (
        "cinematic industrial manga, clean line art, restrained cel shading, "
        "realistic materials, professional adult characters"
    )
    world_lock: str = ""
    palette: list[str] = Field(default_factory=lambda: ["deep navy", "sea teal", "safety orange"])
    prohibited: list[str] = Field(
        default_factory=lambda: [
            "readable generated text",
            "logo",
            "watermark",
            "extra limbs",
            "unsafe action",
        ]
    )
    characters: list[Character] = Field(default_factory=list)
    glossary: list[GlossaryTerm] = Field(default_factory=list)
    created_at: str = Field(default_factory=utc_now)
    updated_at: str = Field(default_factory=utc_now)


class Dialogue(BaseModel):
    dialogue_id: str
    speaker_id: str = "narrator"
    text: str
    bubble_style: str = "rounded_callout"
    layout_mode: Literal["AUTO", "MANUAL"] = "AUTO"
    x: float = 0.08
    y: float = 0.06
    w: float = 0.38
    h: float = 0.12
    font_size: float = 16.0
    tail_target: str = ""
    tail_x: float | None = None
    tail_y: float | None = None

    @field_validator("x", "y", "w", "h", "tail_x", "tail_y")
    @classmethod
    def validate_normalized_coordinate(cls, value: float | None) -> float | None:
        if value is not None and not 0 <= value <= 1:
            raise ValueError("泡泡座標必須介於 0 與 1")
        return value


class ComicPage(BaseModel):
    page_no: int = Field(ge=1, le=80)
    beat: str = ""
    scene_id: str = ""
    scene_description: str = ""
    camera: str = "medium shot"
    learning_point: str = ""
    evidence_ids: list[str] = Field(default_factory=list)
    dialogues: list[Dialogue] = Field(default_factory=list)
    image_prompt: str = ""
    image_asset_id: str | None = None
    alt_text: str = ""


class EvidenceSource(BaseModel):
    source_id: str
    title: str
    source_type: Literal[
        "OEM", "official_guidance", "standard", "research_paper", "course_material", "inference"
    ] = "course_material"
    publisher: str = ""
    url: str = ""
    accessed_at: str = ""
    citation: str = ""
    supported_claims: list[str] = Field(default_factory=list)
    limits: list[str] = Field(default_factory=list)
    page_mapping: list[int] = Field(default_factory=list)


class ComicAsset(BaseModel):
    asset_id: str
    kind: Literal["character_anchor", "equipment_reference", "scene", "draft", "precise_edit"]
    path: str
    sha256: str
    provenance: str
    status: AssetStatus = "DRAFT"
    prompt_version: str = ""
    created_at: str = Field(default_factory=utc_now)


class QARecord(BaseModel):
    gate: str
    result: QAResult
    evidence: str = ""
    reviewer: str = ""
    timestamp: str = Field(default_factory=utc_now)


class ReleaseRecord(BaseModel):
    release_id: str
    channel: Literal["internal_reader", "public_web"] = "internal_reader"
    public_version: str
    published_at: str = Field(default_factory=utc_now)
    published_by: str
    url: str
    withdrawn_at: str | None = None


class EpisodeManifest(BaseModel):
    schema_version: str = "1.0"
    project_id: str
    series_id: str
    story_id: str
    version: str = "v0.1"
    revision: int = 1
    title: str
    week: str = ""
    audience: str = "大學生"
    state: EpisodeState = "BRIEF"
    hold_reason: str = ""
    learning_objectives: list[str] = Field(default_factory=list)
    technical_topics: list[str] = Field(default_factory=list)
    characters: list[str] = Field(default_factory=list)
    page_count: int = Field(default=8, ge=1, le=80)
    story_brief: str = ""
    story_summary: str = ""
    story_beats: list[str] = Field(default_factory=list)
    character_visual_bible: str = ""
    evidence_boundary: str = "teaching_story_not_field_instruction"
    pages: list[ComicPage] = Field(default_factory=list)
    evidence: list[EvidenceSource] = Field(default_factory=list)
    assets: list[ComicAsset] = Field(default_factory=list)
    qa_records: list[QARecord] = Field(default_factory=list)
    exports: dict[str, str] = Field(default_factory=dict)
    releases: list[ReleaseRecord] = Field(default_factory=list)
    external_origin: str = ""
    created_at: str = Field(default_factory=utc_now)
    updated_at: str = Field(default_factory=utc_now)

    @field_validator("version")
    @classmethod
    def validate_version(cls, value: str) -> str:
        if not re.fullmatch(r"v\d+\.\d+", value):
            raise ValueError("version 必須符合 v0.1 格式")
        return value


class ValidationItem(BaseModel):
    check: str
    result: QAResult
    detail: str


class ValidationReport(BaseModel):
    story_id: str
    version: str
    result: QAResult
    publish_ready: bool
    items: list[ValidationItem]
    generated_at: str = Field(default_factory=utc_now)


class PackageDiscovery(BaseModel):
    package_path: str
    story_id: str
    version: str
    title: str
    page_count: int
    scene_count: int
    dialogue_count: int
    source_files: list[str]
    missing_files: list[str]
    historical_qa_claim: str = "UNVERIFIED"
    notes: list[str] = Field(default_factory=list)


class ComicNotFoundError(KeyError):
    pass


class ComicConflictError(RuntimeError):
    pass


class ComicGateError(RuntimeError):
    pass


_PROTECTED_EPISODE_FIELDS = {
    "project_id", "series_id", "story_id", "version", "revision", "created_at", "releases"
}


class ComicStore:
    """漫畫檔案持久層；所有改稿前先保存舊 manifest revision。"""

    def __init__(self, project_root: Path | str = config.PROJECTS_DIR) -> None:
        self.project_root = Path(project_root)
        self._lock = threading.RLock()

    def _comic_root(self, project_id: str) -> Path:
        return self.project_root / safe_id(project_id) / "comics"

    def _series_file(self, project_id: str, series_id: str) -> Path:
        return self._comic_root(project_id) / "series" / safe_id(series_id) / "series.json"

    def episode_dir(self, project_id: str, story_id: str, version: str) -> Path:
        if not re.fullmatch(r"v\d+\.\d+", version):
            raise ValueError("version 必須符合 v0.1 格式")
        return self._comic_root(project_id) / "episodes" / safe_id(story_id) / version

    def _manifest_file(self, project_id: str, story_id: str, version: str) -> Path:
        return self.episode_dir(project_id, story_id, version) / "manifest.json"

    @staticmethod
    def _write_json(path: Path, data: dict[str, Any]) -> None:
        path.parent.mkdir(parents=True, exist_ok=True)
        temp = path.with_suffix(path.suffix + ".tmp")
        temp.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")
        temp.replace(path)

    def create_series(self, series: Series) -> Series:
        clean = series.model_copy(
            update={"project_id": safe_id(series.project_id), "series_id": safe_id(series.series_id)}
        )
        path = self._series_file(clean.project_id, clean.series_id)
        with self._lock:
            if path.is_file():
                raise ComicConflictError(f"series 已存在: {clean.series_id}")
            self._write_json(path, clean.model_dump())
        return clean

    def list_series(self, project_id: str) -> list[Series]:
        root = self._comic_root(project_id) / "series"
        if not root.is_dir():
            return []
        out: list[Series] = []
        for path in sorted(root.glob("*/series.json")):
            try:
                out.append(Series.model_validate_json(path.read_text(encoding="utf-8")))
            except (ValueError, OSError):
                continue
        return out

    def get_series(self, project_id: str, series_id: str) -> Series:
        path = self._series_file(project_id, series_id)
        if not path.is_file():
            raise ComicNotFoundError(f"series 不存在: {series_id}")
        return Series.model_validate_json(path.read_text(encoding="utf-8"))

    def save_series(self, series: Series) -> Series:
        with self._lock:
            self.get_series(series.project_id, series.series_id)
            saved = series.model_copy(update={"updated_at": utc_now()})
            self._write_json(self._series_file(saved.project_id, saved.series_id), saved.model_dump())
            return saved

    def create_episode(self, episode: EpisodeManifest) -> EpisodeManifest:
        clean = episode.model_copy(
            update={
                "project_id": safe_id(episode.project_id),
                "series_id": safe_id(episode.series_id),
                "story_id": safe_id(episode.story_id),
            }
        )
        self.get_series(clean.project_id, clean.series_id)
        path = self._manifest_file(clean.project_id, clean.story_id, clean.version)
        with self._lock:
            if path.is_file():
                raise ComicConflictError(f"episode version 已存在: {clean.story_id}/{clean.version}")
            self._save_episode(clean, archive=False)
        return clean

    def list_episodes(self, project_id: str, series_id: str | None = None) -> list[EpisodeManifest]:
        root = self._comic_root(project_id) / "episodes"
        if not root.is_dir():
            return []
        out: list[EpisodeManifest] = []
        for story_dir in sorted(p for p in root.iterdir() if p.is_dir()):
            versions: list[EpisodeManifest] = []
            for manifest_path in story_dir.glob("*/manifest.json"):
                try:
                    item = EpisodeManifest.model_validate_json(manifest_path.read_text(encoding="utf-8"))
                except (ValueError, OSError):
                    continue
                if series_id is None or item.series_id == safe_id(series_id):
                    versions.append(item)
            if versions:
                versions.sort(key=lambda x: self._version_tuple(x.version), reverse=True)
                out.append(versions[0])
        return out

    def list_versions(self, project_id: str, story_id: str) -> list[EpisodeManifest]:
        root = self._comic_root(project_id) / "episodes" / safe_id(story_id)
        if not root.is_dir():
            return []
        out: list[EpisodeManifest] = []
        for path in root.glob("*/manifest.json"):
            try:
                out.append(EpisodeManifest.model_validate_json(path.read_text(encoding="utf-8")))
            except (ValueError, OSError):
                continue
        return sorted(out, key=lambda item: self._version_tuple(item.version), reverse=True)

    @staticmethod
    def _version_tuple(version: str) -> tuple[int, int]:
        match = re.fullmatch(r"v(\d+)\.(\d+)", version)
        return (int(match.group(1)), int(match.group(2))) if match else (0, 0)

    def get_episode(self, project_id: str, story_id: str, version: str = "v0.1") -> EpisodeManifest:
        path = self._manifest_file(project_id, story_id, version)
        if not path.is_file():
            raise ComicNotFoundError(f"episode 不存在: {story_id}/{version}")
        return EpisodeManifest.model_validate_json(path.read_text(encoding="utf-8"))

    def _archive_manifest(self, episode: EpisodeManifest) -> None:
        history = self.episode_dir(episode.project_id, episode.story_id, episode.version) / "history"
        path = history / f"manifest_r{episode.revision:04d}.json"
        if not path.exists():
            self._write_json(path, episode.model_dump())

    def _save_episode(self, episode: EpisodeManifest, *, archive: bool = True) -> EpisodeManifest:
        path = self._manifest_file(episode.project_id, episode.story_id, episode.version)
        if archive and path.is_file():
            current = EpisodeManifest.model_validate_json(path.read_text(encoding="utf-8"))
            self._archive_manifest(current)
        self._write_json(path, episode.model_dump())
        self._sync_source_pack(episode)
        return episode

    def update_episode(self, project_id: str, story_id: str, version: str, updates: dict[str, Any]) -> EpisodeManifest:
        with self._lock:
            current = self.get_episode(project_id, story_id, version)
            if current.state == "CURRENT":
                raise ComicConflictError("CURRENT 版本不可直接改稿；請建立新版本")
            payload = {k: v for k, v in updates.items() if k not in _PROTECTED_EPISODE_FIELDS}
            payload.update({"revision": current.revision + 1, "updated_at": utc_now()})
            saved = EpisodeManifest.model_validate({**current.model_dump(), **payload})
            return self._save_episode(saved)

    def fork_version(self, project_id: str, story_id: str, from_version: str, new_version: str) -> EpisodeManifest:
        current = self.get_episode(project_id, story_id, from_version)
        if self._manifest_file(project_id, story_id, new_version).is_file():
            raise ComicConflictError(f"episode version 已存在: {story_id}/{new_version}")
        forked = current.model_copy(
            deep=True,
            update={
                "version": new_version,
                "revision": 1,
                "state": "QA" if current.pages else "BRIEF",
                "hold_reason": "",
                "qa_records": [],
                "exports": {},
                "releases": [],
                "created_at": utc_now(),
                "updated_at": utc_now(),
            },
        )
        target_root = self.episode_dir(project_id, story_id, new_version)
        # manifest 內的 asset path 是版本相對路徑；fork 時必須一併複製，不能引用舊版檔案。
        for asset in current.assets:
            source = self.resolve_asset(current, asset.asset_id)
            destination = target_root / asset.path
            destination.parent.mkdir(parents=True, exist_ok=True)
            shutil.copy2(source, destination)
        return self.create_episode(forked)

    def discover_package(self, package_path: Path | str) -> PackageDiscovery:
        root = Path(package_path).expanduser().resolve()
        if not root.is_dir():
            raise ValueError(f"episode package 不存在: {root}")
        match = re.fullmatch(r"(.+)_v(\d+\.\d+)", root.name)
        story_id = match.group(1) if match else root.name
        version = f"v{match.group(2)}" if match else "v0.1"
        source = root / "source"
        required = [
            "storyboard.md",
            "dialogue_script.md",
            "image_prompts.md",
            "technical_sources.md",
            "revision_notes.md",
            "qa_report.md",
        ]
        present = sorted(path.name for path in source.glob("*.md")) if source.is_dir() else []
        missing = [name for name in required if name not in present]
        storyboard_text = self._read_optional(source / "storyboard.md")
        dialogue_text = self._read_optional(source / "dialogue_script.md")
        qa_text = self._read_optional(source / "qa_report.md")
        title_match = re.search(r"《([^》]+)》", storyboard_text or dialogue_text)
        table_rows = self._parse_storyboard_rows(storyboard_text)
        dialogue_blocks = self._parse_dialogue_blocks(dialogue_text)
        scenes = list((root / "assets").glob("scene_*.png")) + list((root / "assets").glob("scene_*.jpg"))
        qa_claim = "PASS" if re.search(r"狀態[：:]\s*PASS", qa_text, re.IGNORECASE) else "UNVERIFIED"
        notes = ["discovery 為唯讀；歷史 QA claim 不會自動轉成新系統 human approval。"]
        if len(table_rows) != len(dialogue_blocks):
            notes.append(f"storyboard 頁數 {len(table_rows)} 與 dialogue 頁數 {len(dialogue_blocks)} 不一致。")
        return PackageDiscovery(
            package_path=str(root),
            story_id=safe_id(story_id),
            version=version,
            title=title_match.group(1) if title_match else story_id,
            page_count=max(len(table_rows), len(dialogue_blocks)),
            scene_count=len(scenes),
            dialogue_count=sum(len(items) for items in dialogue_blocks.values()),
            source_files=present,
            missing_files=missing,
            historical_qa_claim=qa_claim,
            notes=notes,
        )

    def import_package(
        self,
        project_id: str,
        series_id: str,
        package_path: Path | str,
    ) -> EpisodeManifest:
        """唯讀解析既有 package，複製資產到新 Project；不修改原始資料夾。"""
        report = self.discover_package(package_path)
        if report.page_count < 1:
            raise ComicGateError("找不到 storyboard/dialogue 頁面，未匯入")
        if report.missing_files:
            raise ComicGateError(f"缺少必要 source files: {', '.join(report.missing_files)}")
        self.get_series(project_id, series_id)
        origin = Path(report.package_path)
        source = origin / "source"
        storyboard_rows = self._parse_storyboard_rows(self._read_optional(source / "storyboard.md"))
        dialogue_blocks = self._parse_dialogue_blocks(self._read_optional(source / "dialogue_script.md"))
        prompt_map, scene_page_map = self._parse_prompt_sections(self._read_optional(source / "image_prompts.md"))
        evidence = self._parse_imported_evidence(self._read_optional(source / "technical_sources.md"))
        objective_match = re.search(
            r"Learning objective[：:]\s*(.+)",
            self._read_optional(source / "storyboard.md"),
            re.IGNORECASE,
        )
        pages: list[ComicPage] = []
        for page_no in range(1, report.page_count + 1):
            row = storyboard_rows.get(page_no, {})
            dialogues = [
                Dialogue(
                    dialogue_id=f"p{page_no:02d}_d{idx:02d}",
                    speaker_id=self._speaker_id(speaker),
                    text=text,
                )
                for idx, (speaker, text) in enumerate(dialogue_blocks.get(page_no, []), start=1)
            ]
            pages.append(
                ComicPage(
                    page_no=page_no,
                    beat=row.get("title", ""),
                    scene_id=f"scene_{page_no:02d}",
                    scene_description=row.get("scene", row.get("title", "")),
                    learning_point=row.get("learning_point", ""),
                    dialogues=dialogues,
                    image_prompt=prompt_map.get(page_no, ""),
                    alt_text=f"{report.title}第 {page_no} 頁：{row.get('scene', row.get('title', '既有漫畫場景'))}",
                )
            )
        episode = self.create_episode(
            EpisodeManifest(
                project_id=project_id,
                series_id=series_id,
                story_id=report.story_id,
                version=report.version,
                title=report.title,
                week=report.story_id.split("_", 1)[0],
                state="STORYBOARD",
                page_count=report.page_count,
                story_brief="由既有 episode package 唯讀匯入；內容與技術結論仍須在新系統重新確認。",
                learning_objectives=[objective_match.group(1).strip()] if objective_match else ["待教師確認既有 learning objective"],
                characters=sorted({item.speaker_id for page in pages for item in page.dialogues}),
                pages=pages,
                evidence=evidence,
                external_origin=str(origin),
            )
        )
        scene_files = sorted(
            [*(origin / "assets").glob("scene_*.png"), *(origin / "assets").glob("scene_*.jpg")]
        )
        scene_assets: dict[int, str] = {}
        for fallback_index, scene_path in enumerate(scene_files, start=1):
            number_match = re.match(r"scene_(\d+)", scene_path.stem)
            scene_no = int(number_match.group(1)) if number_match else fallback_index
            asset_id = f"import_scene_{scene_no:02d}"
            episode = self.attach_asset(
                project_id,
                report.story_id,
                report.version,
                filename=scene_path.name,
                data=scene_path.read_bytes(),
                kind="scene",
                provenance=f"imported:{origin}",
                asset_id=asset_id,
                status="FINAL",
                prompt_version="historical_import",
            )
            scene_assets[scene_no] = asset_id
        for page_no in range(1, report.page_count + 1):
            scene_no = scene_page_map.get(page_no)
            if scene_no in scene_assets:
                episode = self.set_page_asset(project_id, report.story_id, report.version, page_no, scene_assets[scene_no])
        # 保留原始 Markdown/建置腳本的唯讀副本；新系統產生的 source/*.md 仍為 normalized view。
        original_copy = self.episode_dir(project_id, report.story_id, report.version) / "source" / "original_import"
        original_copy.mkdir(parents=True, exist_ok=True)
        for path in source.iterdir():
            if path.is_file() and path.suffix.lower() in {".md", ".py", ".ps1"}:
                shutil.copy2(path, original_copy / path.name)
        return self.update_episode(
            project_id,
            report.story_id,
            report.version,
            {
                "state": "HOLD",
                "hold_reason": "歷史 QA 已偵測，但新系統 human approval 尚未完成。",
            },
        )

    @staticmethod
    def _read_optional(path: Path) -> str:
        return path.read_text(encoding="utf-8-sig") if path.is_file() else ""

    @staticmethod
    def _parse_storyboard_rows(text: str) -> dict[int, dict[str, str]]:
        rows: dict[int, dict[str, str]] = {}
        for match in re.finditer(r"^\|\s*(\d{1,2})\s*\|\s*([^|]+)\|\s*([^|]+)\|\s*([^|]+)\|", text, re.MULTILINE):
            page_no = int(match.group(1))
            title_scene = match.group(2).strip()
            title, _, scene = title_scene.partition("／")
            rows[page_no] = {
                "title": title.strip(),
                "scene": (scene or title).strip(),
                "learning_point": match.group(3).strip(),
                "composition": match.group(4).strip(),
            }
        return rows

    @staticmethod
    def _parse_dialogue_blocks(text: str) -> dict[int, list[tuple[str, str]]]:
        blocks: dict[int, list[tuple[str, str]]] = {}
        matches = list(re.finditer(r"^##\s*P(\d{1,2})[^\n]*$", text, re.MULTILINE))
        for index, match in enumerate(matches):
            page_no = int(match.group(1))
            end = matches[index + 1].start() if index + 1 < len(matches) else len(text)
            body = text[match.end():end]
            blocks[page_no] = [
                (item.group(1).strip(), item.group(2).strip())
                for item in re.finditer(r"^-\s*([^：:\n]+)[：:]\s*(.+)$", body, re.MULTILINE)
            ]
        return blocks

    @staticmethod
    def _parse_prompt_sections(text: str) -> tuple[dict[int, str], dict[int, int]]:
        prompts: dict[int, str] = {}
        mapping: dict[int, int] = {}
        matches = list(re.finditer(r"^##\s*Scene\s*(\d+)[^\n]*\(P(\d+)(?:[–-]P?(\d+))?\)[^\n]*$", text, re.MULTILINE))
        continuity = text.split("## Scene", 1)[0].strip()
        for index, match in enumerate(matches):
            scene_no = int(match.group(1))
            first = int(match.group(2))
            last = int(match.group(3) or first)
            end = matches[index + 1].start() if index + 1 < len(matches) else len(text)
            prompt = continuity + "\n\n" + text[match.start():end].strip()
            for page_no in range(first, last + 1):
                prompts[page_no] = prompt
                mapping[page_no] = scene_no
        return prompts, mapping

    @staticmethod
    def _parse_imported_evidence(text: str) -> list[EvidenceSource]:
        items: list[EvidenceSource] = []
        source_part = text.split("## INFERENCE", 1)[0]
        matches = list(re.finditer(r"^\d+\.\s+(.+?)(?:\s{2,}|\n)\s*(https?://\S+)", source_part, re.MULTILINE | re.DOTALL))
        for index, match in enumerate(matches[:30], start=1):
            description = " ".join(match.group(1).split())
            items.append(
                EvidenceSource(
                    source_id=f"import_src_{index:02d}",
                    title=description[:180],
                    source_type="course_material",
                    url=match.group(2).rstrip(".)"),
                    citation=description,
                    supported_claims=["由既有 technical_sources.md 匯入；需 technical reviewer 對頁面映射重驗。"],
                    limits=["歷史來源匯入不等於新系統已完成 claim-level verification。"],
                )
            )
        return items

    @staticmethod
    def _speaker_id(name: str) -> str:
        known = {"杜夫": "dofu", "小櫻": "sakura", "若晴": "ruoqing", "林若晴": "ruoqing", "陳教練": "instructor_chen"}
        return known.get(name, f"speaker_{hashlib.sha1(name.encode('utf-8')).hexdigest()[:8]}")

    def add_evidence(self, project_id: str, story_id: str, version: str, evidence: EvidenceSource) -> EpisodeManifest:
        current = self.get_episode(project_id, story_id, version)
        items = [item for item in current.evidence if item.source_id != evidence.source_id]
        items.append(evidence)
        return self.update_episode(project_id, story_id, version, {"evidence": items})

    def add_qa_record(self, project_id: str, story_id: str, version: str, record: QARecord) -> EpisodeManifest:
        current = self.get_episode(project_id, story_id, version)
        records = [item for item in current.qa_records if item.gate != record.gate]
        records.append(record)
        return self.update_episode(project_id, story_id, version, {"qa_records": records})

    def transition(self, project_id: str, story_id: str, version: str, target: EpisodeState, reason: str = "") -> EpisodeManifest:
        current = self.get_episode(project_id, story_id, version)
        if target == "HOLD":
            if not reason.strip():
                raise ComicGateError("HOLD 必須填寫原因")
            return self.update_episode(project_id, story_id, version, {"state": "HOLD", "hold_reason": reason})
        if current.state == "CURRENT":
            raise ComicConflictError("CURRENT 版本不可變更狀態；請 fork 新版本")
        if target == "CURRENT":
            report = self.validate_episode(current)
            if not report.publish_ready:
                raise ComicGateError("尚未通過全部發布 gate")
        if target not in STATE_ORDER:
            raise ComicGateError(f"不支援的狀態: {target}")
        return self.update_episode(project_id, story_id, version, {"state": target, "hold_reason": ""})

    def attach_asset(
        self,
        project_id: str,
        story_id: str,
        version: str,
        *,
        filename: str,
        data: bytes,
        kind: str,
        provenance: str,
        asset_id: str | None = None,
        status: AssetStatus = "DRAFT",
        prompt_version: str = "",
    ) -> EpisodeManifest:
        if len(data) > 25 * 1024 * 1024:
            raise ValueError("單一 asset 不可超過 25 MB")
        suffix = Path(filename).suffix.lower()
        if suffix not in {".png", ".jpg", ".jpeg", ".webp"}:
            raise ValueError("asset 僅支援 PNG/JPG/WEBP")
        current = self.get_episode(project_id, story_id, version)
        aid = safe_id(asset_id or f"asset_{uuid.uuid4().hex[:8]}")
        rel = Path("assets") / f"{aid}{suffix}"
        out = self.episode_dir(project_id, story_id, version) / rel
        out.parent.mkdir(parents=True, exist_ok=True)
        if out.exists():
            raise ComicConflictError(f"asset 已存在: {aid}")
        out.write_bytes(data)
        asset = ComicAsset(
            asset_id=aid,
            kind=kind,
            path=rel.as_posix(),
            sha256=sha256_file(out),
            provenance=provenance,
            status=status,
            prompt_version=prompt_version,
        )
        assets = [item for item in current.assets if item.asset_id != aid] + [asset]
        return self.update_episode(project_id, story_id, version, {"assets": assets})

    def set_page_asset(self, project_id: str, story_id: str, version: str, page_no: int, asset_id: str) -> EpisodeManifest:
        current = self.get_episode(project_id, story_id, version)
        if not any(item.asset_id == asset_id for item in current.assets):
            raise ComicNotFoundError(f"asset 不存在: {asset_id}")
        pages = [page.model_copy(update={"image_asset_id": asset_id}) if page.page_no == page_no else page for page in current.pages]
        if not any(page.page_no == page_no for page in current.pages):
            raise ComicNotFoundError(f"page 不存在: {page_no}")
        target = "IMAGE" if pages and all(page.image_asset_id for page in pages) else current.state
        return self.update_episode(project_id, story_id, version, {"pages": pages, "state": target})

    def resolve_asset(self, episode: EpisodeManifest, asset_id: str) -> Path:
        asset = next((item for item in episode.assets if item.asset_id == asset_id), None)
        if asset is None:
            raise ComicNotFoundError(f"asset 不存在: {asset_id}")
        root = self.episode_dir(episode.project_id, episode.story_id, episode.version).resolve()
        path = (root / asset.path).resolve()
        try:
            path.relative_to(root)
        except ValueError as exc:
            raise ComicGateError("asset path 超出 episode package") from exc
        if not path.is_file():
            raise ComicNotFoundError(f"asset 檔案不存在: {asset.path}")
        return path

    @staticmethod
    def _rect_overlap(a: tuple[float, float, float, float], b: tuple[float, float, float, float]) -> float:
        left = max(a[0], b[0])
        top = max(a[1], b[1])
        right = min(a[0] + a[2], b[0] + b[2])
        bottom = min(a[1] + a[3], b[1] + b[3])
        return max(0.0, right - left) * max(0.0, bottom - top)

    def resolve_dialogue_layout(self, episode: EpisodeManifest, page: ComicPage) -> list[Dialogue]:
        """依實際場景圖的低細節區配置泡泡；MANUAL 座標則完整保留。"""
        if not page.dialogues:
            return []

        try:
            series = self.get_series(episode.project_id, episode.series_id)
            character_ids = [character.character_id for character in series.characters]
        except ComicNotFoundError:
            character_ids = list(dict.fromkeys(episode.characters))
        if len(character_ids) == 1:
            speaker_anchors = {character_ids[0]: 0.50}
        elif character_ids:
            speaker_anchors = {
                character_id: 0.34 + index * 0.36 / (len(character_ids) - 1)
                for index, character_id in enumerate(character_ids)
            }
        else:
            speaker_anchors = {}
        speaker_target_y = {character_id: 0.72 for character_id in character_ids}

        edge_map = None
        if page.image_asset_id:
            try:
                from PIL import Image, ImageFilter

                asset_path = self.resolve_asset(episode, page.image_asset_id)
                with Image.open(asset_path) as source:
                    edge_map = source.convert("L").resize((240, 320)).filter(ImageFilter.FIND_EDGES)
            except (OSError, ComicNotFoundError):
                edge_map = None
            else:
                # 本機 OpenCV 只用來找人臉位置，不做人名辨識；角色仍依 Series 順序由左至右配對。
                try:
                    import cv2

                    frame = cv2.imread(str(asset_path))
                    if frame is not None and character_ids:
                        gray = cv2.cvtColor(frame, cv2.COLOR_BGR2GRAY)
                        detected: list[tuple[float, float, float, float]] = []
                        for cascade_name, scale, neighbors, flip, confidence in (
                            ("haarcascade_frontalface_default.xml", 1.08, 4, False, 1.0),
                            ("haarcascade_frontalface_default.xml", 1.05, 3, False, 0.9),
                            ("haarcascade_profileface.xml", 1.05, 3, False, 0.65),
                            ("haarcascade_profileface.xml", 1.05, 3, True, 0.65),
                        ):
                            cascade = cv2.CascadeClassifier(cv2.data.haarcascades + cascade_name)
                            detection_frame = cv2.flip(gray, 1) if flip else gray
                            for face_x, face_y, face_w, face_h in cascade.detectMultiScale(
                                detection_frame, scaleFactor=scale, minNeighbors=neighbors, minSize=(36, 36)
                            ):
                                center_x = (face_x + face_w / 2) / frame.shape[1]
                                if flip:
                                    center_x = 1 - center_x
                                center_y = (face_y + face_h / 2) / frame.shape[0]
                                width_ratio = face_w / frame.shape[1]
                                if 0.28 <= center_y <= 0.72 and 0.035 <= width_ratio <= 0.28:
                                    detected.append((center_x, center_y, width_ratio, width_ratio * confidence))
                        unique_faces: list[tuple[float, float, float, float]] = []
                        for candidate in sorted(detected, key=lambda item: item[3], reverse=True):
                            if all(
                                abs(candidate[0] - prior[0]) > 0.08 or abs(candidate[1] - prior[1]) > 0.08
                                for prior in unique_faces
                            ):
                                unique_faces.append(candidate)
                        if len(unique_faces) >= len(character_ids):
                            visible = sorted(unique_faces[:len(character_ids)], key=lambda item: item[0])
                            speaker_anchors = {
                                character_id: round(face[0], 4)
                                for character_id, face in zip(character_ids, visible)
                            }
                            speaker_target_y = {
                                character_id: round(min(0.84, face[1] + 0.12), 4)
                                for character_id, face in zip(character_ids, visible)
                            }
                except (ImportError, OSError):
                    pass

        # 候選點分散在不同高度；依頁碼輪替同分候選，避免整個 episode 都排成同一列。
        candidates = [
            (0.05, 0.06), (0.55, 0.06),
            (0.05, 0.25), (0.55, 0.25),
            (0.05, 0.44), (0.55, 0.44),
            (0.10, 0.62), (0.52, 0.62),
        ]
        offset = (page.page_no * 3) % len(candidates)
        candidates = candidates[offset:] + candidates[:offset]
        occupied: list[tuple[float, float, float, float]] = []
        resolved: list[Dialogue] = []

        layout_dialogues = page.dialogues[:3]
        overflow_dialogues = page.dialogues[3:]

        for index, dialogue in enumerate(layout_dialogues):
            text_units = sum(2 if ord(char) > 127 else 1 for char in dialogue.text)
            if dialogue.layout_mode == "MANUAL":
                width = min(0.70, max(0.18, dialogue.w))
                height = min(0.30, max(0.08, dialogue.h))
                x = min(max(0.02, dialogue.x), 0.98 - width)
                y = min(max(0.02, dialogue.y), 0.82 - height)
            else:
                width = min(0.46, max(0.32, dialogue.w, 0.34 + max(0, text_units - 28) * 0.002))
                line_capacity = max(12, int(width * 70))
                estimated_lines = max(2, (text_units + line_capacity - 1) // line_capacity)
                height = min(0.20, max(0.115, dialogue.h, 0.075 + estimated_lines * 0.026))
                ranked: list[tuple[float, float, float]] = []
                for order, (candidate_x, candidate_y) in enumerate(candidates):
                    x0 = min(candidate_x, 0.95 - width)
                    y0 = min(candidate_y, 0.82 - height)
                    rect = (x0, y0, width, height)
                    overlap = sum(self._rect_overlap(rect, prior) for prior in occupied)
                    detail = 0.0
                    if edge_map is not None:
                        left = max(0, int(x0 * edge_map.width))
                        top = max(0, int(y0 * edge_map.height))
                        right = min(edge_map.width, max(left + 1, int((x0 + width) * edge_map.width)))
                        bottom = min(edge_map.height, max(top + 1, int((y0 + height) * edge_map.height)))
                        histogram = edge_map.crop((left, top, right, bottom)).histogram()
                        pixel_count = max(1, sum(histogram))
                        detail = sum(level * count for level, count in enumerate(histogram)) / (255 * pixel_count)
                    speaker_x = speaker_anchors.get(dialogue.speaker_id, 0.5)
                    proximity = abs((x0 + width / 2) - speaker_x)
                    tail_gap = max(0.0, 0.72 - (y0 + height))
                    tail_ergonomics = abs(tail_gap - 0.24)
                    # overlap 是硬性高權重；order 只負責穩定打破平手。
                    ranked.append((detail + proximity * 0.12 + tail_ergonomics * 0.18 + overlap * 80 + order * 0.0005, x0, y0))
                _, x, y = min(ranked, key=lambda item: item[0])

            manual_tail = dialogue.layout_mode == "MANUAL"
            target_x = dialogue.tail_x if manual_tail and dialogue.tail_x is not None else speaker_anchors.get(dialogue.speaker_id, 0.5)
            target_y = (
                dialogue.tail_y
                if manual_tail and dialogue.tail_y is not None
                else max(speaker_target_y.get(dialogue.speaker_id, 0.72), y + height + 0.08)
            )
            target_x = min(0.94, max(0.06, target_x))
            target_y = min(0.94, max(y + height + 0.045, target_y))
            occupied.append((x, y, width, height))
            resolved.append(
                dialogue.model_copy(
                    update={
                        "bubble_style": "rounded_callout",
                        "x": round(x, 4),
                        "y": round(y, 4),
                        "w": round(width, 4),
                        "h": round(height, 4),
                        "tail_x": round(target_x, 4),
                        "tail_y": round(target_y, 4),
                    }
                )
            )
        return [*resolved, *overflow_dialogues]

    def auto_layout_episode(self, project_id: str, story_id: str, version: str) -> EpisodeManifest:
        """把 image-aware 配置寫回 manifest，供 UI 預覽與後續人工微調。"""
        episode = self.get_episode(project_id, story_id, version)
        pages = [
            page.model_copy(update={"dialogues": self.resolve_dialogue_layout(episode, page)})
            for page in episode.pages
        ]
        target_state = episode.state
        if target_state in {"BRIEF", "STORYBOARD", "IMAGE"} and pages:
            target_state = "LAYOUT"
        return self.update_episode(project_id, story_id, version, {"pages": pages, "state": target_state})

    def compose_prompts(self, project_id: str, story_id: str, version: str) -> EpisodeManifest:
        episode = self.get_episode(project_id, story_id, version)
        series = self.get_series(project_id, episode.series_id)
        char_map = {item.character_id: item for item in series.characters}
        char_lock = "\n".join(
            f"- {cid}: {char_map[cid].visual_lock}" for cid in episode.characters if cid in char_map
        ) or episode.character_visual_bible or "Preserve the same named adult characters and wardrobe."
        negative = ", ".join(series.prohibited)
        pages: list[ComicPage] = []
        for page in episode.pages:
            prompt = f"""USE CASE: illustration-story.
Create one vertical A4 educational manga background for the same serialized comic.

CHARACTER LOCK — do not redesign:
{char_lock}

WORLD AND STYLE LOCK:
{series.visual_bible}. {series.world_lock}

SCENE:
{page.scene_description}
Camera: {page.camera}. Visible learning evidence: {page.learning_point}.

COMPOSITION:
Place people, hands, PPE, tools and technical evidence in a 62–66% action zone. Preserve 34–38% integrated negative space for later Traditional Chinese speech bubbles.

CONSTRAINTS:
No generated captions, no readable Chinese text, no speech bubbles. Teaching story, not field operating instruction.
Negative prompt: {negative}.
""".strip()
            pages.append(page.model_copy(update={"image_prompt": prompt}))
        return self.update_episode(project_id, story_id, version, {"pages": pages, "state": "STORYBOARD"})

    def validate_episode(self, episode: EpisodeManifest) -> ValidationReport:
        items: list[ValidationItem] = []

        def add(check: str, passed: bool, detail: str, *, hold: bool = False) -> None:
            result: QAResult = "PASS" if passed else ("HOLD" if hold else "FAIL")
            items.append(ValidationItem(check=check, result=result, detail=detail))

        add("identity", bool(episode.story_id and episode.title and episode.version), "Story ID、標題與版本必填")
        add("learning_objectives", bool(episode.learning_objectives), "至少一項 learning objective")
        add("evidence_boundary", bool(episode.evidence_boundary), "需聲明 teaching/evidence boundary")
        add("evidence_pack", bool(episode.evidence), "至少一筆可追溯來源", hold=not episode.evidence)
        add(
            "pages",
            len(episode.pages) == episode.page_count and [p.page_no for p in episode.pages] == list(range(1, episode.page_count + 1)),
            f"預期 {episode.page_count} 頁，實際 {len(episode.pages)} 頁且頁碼需連續",
        )
        add("dialogue", bool(episode.pages) and all(page.dialogues for page in episode.pages), "每頁至少一段可編輯對白")
        add("prompts", bool(episode.pages) and all(page.image_prompt for page in episode.pages), "每頁需有可重現 image prompt")
        add("scene_assets", bool(episode.pages) and all(page.image_asset_id for page in episode.pages), "每頁需連結 scene asset")
        asset_map = {asset.asset_id: asset for asset in episode.assets}
        used_assets = [asset_map.get(page.image_asset_id or "") for page in episode.pages]
        add(
            "asset_provenance",
            bool(used_assets)
            and all(asset is not None and asset.provenance != "mock_placeholder" for asset in used_assets),
            "正式發布不可使用 mock placeholder；每張 scene asset 必須保留 provenance",
        )
        add("alt_text", bool(episode.pages) and all(page.alt_text for page in episode.pages), "每頁需有 reader alt text")

        latest = {record.gate: record for record in episode.qa_records}
        for gate in REQUIRED_QA_GATES:
            record = latest.get(gate)
            add(f"qa:{gate}", bool(record and record.result == "PASS"), f"{gate} 必須由 reviewer 記錄 PASS", hold=record is None)

        hard_fail = any(item.result == "FAIL" for item in items)
        has_hold = any(item.result in {"HOLD", "UNVERIFIED"} for item in items)
        result: QAResult = "FAIL" if hard_fail else ("HOLD" if has_hold else "PASS")
        return ValidationReport(
            story_id=episode.story_id,
            version=episode.version,
            result=result,
            publish_ready=result == "PASS",
            items=items,
        )

    def build_reader_html(
        self,
        episode: EpisodeManifest,
        *,
        asset_prefix: str = "",
        speaker_names: dict[str, str] | None = None,
    ) -> str:
        speaker_names = speaker_names or {}
        pages_html: list[str] = []
        for page in episode.pages:
            image = ""
            if page.image_asset_id:
                src = f"{asset_prefix}{html.escape(page.image_asset_id)}"
                image = f'<img src="{src}" alt="{html.escape(page.alt_text)}" loading="lazy">'
            dialogue = "".join(
                f'<p><strong>{html.escape(speaker_names.get(item.speaker_id, item.speaker_id))}</strong>：{html.escape(item.text)}</p>'
                for item in page.dialogues
            )
            pages_html.append(
                f'<article class="page" id="page-{page.page_no}">{image}'
                f'<section class="transcript"><span>PAGE {page.page_no:02d}</span>{dialogue}</section></article>'
            )
        objectives = "".join(f"<li>{html.escape(item)}</li>" for item in episode.learning_objectives)
        return f"""<!doctype html>
<html lang="zh-TW"><head><meta charset="utf-8"><meta name="viewport" content="width=device-width,initial-scale=1">
<title>{html.escape(episode.title)}</title><style>
:root{{color-scheme:dark;--bg:#0c1117;--panel:#151d26;--text:#e9f1f7;--muted:#a6b6c5;--accent:#42c7a5}}
*{{box-sizing:border-box}}body{{margin:0;background:var(--bg);color:var(--text);font:16px/1.7 system-ui,"Microsoft JhengHei",sans-serif}}
header,main,footer{{width:min(920px,calc(100% - 28px));margin:auto}}header{{padding:38px 0 24px}}h1{{margin:.15em 0;font-size:clamp(30px,7vw,52px)}}
.boundary{{border-left:4px solid #f3a847;background:#2a2115;padding:12px 16px;color:#ffdca7}}.page{{margin:24px 0;background:var(--panel);border:1px solid #2a3947;border-radius:18px;overflow:hidden}}
.page img{{display:block;width:100%;height:auto}}.transcript{{padding:16px 20px}}.transcript>span{{color:var(--accent);font-size:12px;letter-spacing:.12em}}.transcript p{{margin:.45em 0}}
footer{{padding:20px 0 60px;color:var(--muted)}}a{{color:var(--accent)}}
</style></head><body><header><small>{html.escape(episode.story_id)} · {html.escape(episode.version)}</small><h1>{html.escape(episode.title)}</h1>
<p>{html.escape(episode.story_summary)}</p><div class="boundary">Teaching story；技術與安全處置仍以 OEM、site procedure 與正式授權為準。</div>
<h2>學習目標</h2><ul>{objectives}</ul></header><main>{''.join(pages_html)}</main>
<footer>eduStudio Serialized Comic · Versioned release · {html.escape(episode.updated_at)}</footer></body></html>"""

    def build_series_archive_html(self, series: Series, episodes: list[EpisodeManifest]) -> str:
        episode_cards: list[str] = []
        for episode in episodes:
            active = [release for release in episode.releases if release.withdrawn_at is None]
            if not active:
                continue
            episode_cards.append(
                f'<a class="episode" href="/projects/{html.escape(series.project_id)}/comics/reader/{html.escape(episode.story_id)}">'
                f'<small>{html.escape(episode.week or episode.story_id)} · {html.escape(episode.version)}</small>'
                f'<h2>{html.escape(episode.title)}</h2><p>{html.escape(episode.story_summary)}</p>'
                f'<span>{episode.page_count} pages · {len(episode.evidence)} sources</span></a>'
            )
        character_cards = "".join(
            f'<article><h3>{html.escape(item.name)}</h3><strong>{html.escape(item.role)}</strong><p>{html.escape(item.voice)}</p></article>'
            for item in series.characters
        )
        glossary = "".join(
            f'<dt>{html.escape(item.term)}</dt><dd>{html.escape(item.definition)}</dd>'
            for item in series.glossary
        )
        return f"""<!doctype html><html lang="zh-TW"><head><meta charset="utf-8"><meta name="viewport" content="width=device-width,initial-scale=1">
<title>{html.escape(series.title)}</title><style>
:root{{color-scheme:dark;--bg:#0c1117;--panel:#151d26;--text:#edf5f9;--muted:#9cb0bf;--accent:#43c7a4}}
*{{box-sizing:border-box}}body{{margin:0;background:var(--bg);color:var(--text);font:16px/1.7 system-ui,"Microsoft JhengHei",sans-serif}}
main{{width:min(1040px,calc(100% - 28px));margin:auto;padding:42px 0 70px}}h1{{font-size:clamp(36px,8vw,64px);margin:.1em 0}}.lead{{color:var(--muted);max-width:720px}}
.episodes,.characters{{display:grid;grid-template-columns:repeat(auto-fit,minmax(240px,1fr));gap:14px}}.episode,.characters article{{display:block;padding:20px;border:1px solid #2a3947;border-radius:16px;background:var(--panel);color:var(--text);text-decoration:none}}
.episode:hover{{border-color:var(--accent)}}.episode small,.episode span,.characters strong{{color:var(--accent)}}.episode p,.characters p,dd{{color:var(--muted)}}section{{margin-top:38px}}dl{{display:grid;grid-template-columns:minmax(110px,.25fr) 1fr;gap:8px 16px}}dt{{font-weight:700;color:var(--accent)}}dd{{margin:0}}
</style></head><body><main><header><small>SERIALIZED COMIC ARCHIVE</small><h1>{html.escape(series.title)}</h1><p class="lead">{html.escape(series.description)}</p></header>
<section><h2>已發布 Episodes</h2><div class="episodes">{''.join(episode_cards) or '<p>尚無已發布 episode。</p>'}</div></section>
<section><h2>角色</h2><div class="characters">{character_cards or '<p>尚無角色資料。</p>'}</div></section>
<section><h2>Glossary</h2><dl>{glossary or '<p>尚無術語。</p>'}</dl></section></main></body></html>"""

    def export_html(self, project_id: str, story_id: str, version: str) -> tuple[EpisodeManifest, Path]:
        episode = self.get_episode(project_id, story_id, version)
        series = self.get_series(project_id, episode.series_id)
        speaker_names = {item.character_id: item.name for item in series.characters}
        out = self.episode_dir(project_id, story_id, version) / "exports" / f"{episode.story_id}_{version}_reader.html"
        out.parent.mkdir(parents=True, exist_ok=True)
        prefix = f"/projects/{episode.project_id}/comics/episodes/{episode.story_id}/{version}/assets/"
        out.write_text(
            self.build_reader_html(episode, asset_prefix=prefix, speaker_names=speaker_names),
            encoding="utf-8",
        )
        updated = self._record_exports(
            episode,
            {"html": out.relative_to(self.episode_dir(project_id, story_id, version)).as_posix()},
        )
        return updated, out

    def export_pdf(self, project_id: str, story_id: str, version: str) -> tuple[EpisodeManifest, Path]:
        from PIL import Image, ImageDraw, ImageFont

        episode = self.get_episode(project_id, story_id, version)
        if not episode.pages or any(not page.image_asset_id for page in episode.pages):
            raise ComicGateError("PDF 匯出需要每頁 scene asset")
        series = self.get_series(project_id, episode.series_id)
        speaker_names = {item.character_id: item.name for item in series.characters}
        font_path = config.get_font_path()
        font = ImageFont.truetype(font_path, 30)
        small = ImageFont.truetype(font_path, 21)

        def wrap_text(value: str, target_font: Any, max_width: int) -> list[str]:
            lines: list[str] = []
            current = ""
            for char in value:
                candidate = current + char
                if current and target_font.getlength(candidate) > max_width:
                    lines.append(current)
                    current = char
                else:
                    current = candidate
            if current:
                lines.append(current)
            return lines

        rendered: list[Image.Image] = []
        for page in episode.pages:
            source = Image.open(self.resolve_asset(episode, page.image_asset_id or "")).convert("RGB")
            canvas = Image.new("RGB", (1240, 1754), "white")
            source.thumbnail((1160, 1390))
            x = (1240 - source.width) // 2
            image_top = 70
            canvas.paste(source, (x, image_top))
            draw = ImageDraw.Draw(canvas)
            for dialog in self.resolve_dialogue_layout(episode, page):
                left = int(x + dialog.x * source.width)
                top = int(image_top + dialog.y * source.height)
                right = int(left + dialog.w * source.width)
                bottom = int(top + dialog.h * source.height)
                target_x = int(x + (dialog.tail_x or dialog.x + dialog.w / 2) * source.width)
                target_y = int(image_top + (dialog.tail_y or dialog.y + dialog.h + 0.1) * source.height)
                base_x = min(right - 34, max(left + 34, target_x))
                fill = "#ffffff"
                outline = "#173f5f"
                # 三角尾巴先畫、泡泡本體後畫，交界被本體覆蓋成一體成形。
                tail_points = [(base_x - 24, bottom - 3), (base_x + 24, bottom - 3), (target_x, target_y)]
                draw.polygon(tail_points, fill=fill, outline=outline)
                draw.line([tail_points[0], tail_points[2], tail_points[1]], fill=outline, width=4, joint="curve")
                draw.rounded_rectangle((left, top, right, bottom), radius=28, fill=fill, outline=outline, width=4)
                text = f"{speaker_names.get(dialog.speaker_id, dialog.speaker_id)}：{dialog.text}"
                lines = wrap_text(text, small, max(80, right - left - 44))
                line_height = small.getbbox("國Ay")[3] + 5
                text_height = len(lines) * line_height
                text_y = max(top + 15, top + (bottom - top - text_height) // 2)
                for line in lines[:4]:
                    draw.text((left + 22, text_y), line, font=small, fill="#10202c")
                    text_y += line_height
            draw.text((60, 28), f"{episode.story_id} · {episode.version} · P{page.page_no:02d}", font=font, fill="#173042")
            rendered.append(canvas)
        out = self.episode_dir(project_id, story_id, version) / "exports" / f"{episode.story_id}_{version}.pdf"
        out.parent.mkdir(parents=True, exist_ok=True)
        rendered[0].save(out, "PDF", resolution=150.0, save_all=True, append_images=rendered[1:])
        updated = self._record_exports(
            episode,
            {"pdf": out.relative_to(self.episode_dir(project_id, story_id, version)).as_posix()},
        )
        return updated, out

    def export_docx(self, project_id: str, story_id: str, version: str) -> tuple[EpisodeManifest, Path, str]:
        """先嘗試 Word native Shapes；無 Word COM 時回退為可編輯 table draft。"""
        episode = self.get_episode(project_id, story_id, version)
        if not episode.pages or any(not page.image_asset_id for page in episode.pages):
            raise ComicGateError("DOCX 匯出需要每頁 scene asset")
        out = self.episode_dir(project_id, story_id, version) / "exports" / f"{episode.story_id}_{version}_editable.docx"
        out.parent.mkdir(parents=True, exist_ok=True)
        mode = "word_native_shapes"
        try:
            self._export_docx_word_shapes(episode, out)
        except (ImportError, OSError, RuntimeError):
            mode = "editable_table_fallback"
            self._export_docx_fallback(episode, out)
        current = self.get_episode(project_id, story_id, version)
        updated = self._record_exports(
            current,
            {
                "docx": out.relative_to(self.episode_dir(project_id, story_id, version)).as_posix(),
                "docx_mode": mode,
            },
        )
        return updated, out, mode

    def _record_exports(self, episode: EpisodeManifest, additions: dict[str, str]) -> EpisodeManifest:
        """匯出不改故事內容；CURRENT 亦允許追加 artifact audit，不要求 fork。"""
        exports = {**episode.exports, **additions}
        if episode.state != "CURRENT":
            return self.update_episode(
                episode.project_id,
                episode.story_id,
                episode.version,
                {"exports": exports},
            )
        with self._lock:
            saved = episode.model_copy(
                update={
                    "exports": exports,
                    "revision": episode.revision + 1,
                    "updated_at": utc_now(),
                }
            )
            return self._save_episode(saved)

    def _export_docx_fallback(self, episode: EpisodeManifest, out: Path) -> None:
        from docx import Document
        from docx.enum.text import WD_BREAK
        from docx.shared import Cm, Pt

        doc = Document()
        section = doc.sections[0]
        section.page_width, section.page_height = Cm(21), Cm(29.7)
        section.top_margin = section.bottom_margin = Cm(1.2)
        section.left_margin = section.right_margin = Cm(1.2)
        for idx, page in enumerate(episode.pages):
            if idx:
                doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
            doc.add_heading(f"{episode.title} · P{page.page_no:02d}", level=1)
            doc.add_picture(str(self.resolve_asset(episode, page.image_asset_id or "")), width=Cm(18.2))
            table = doc.add_table(rows=max(1, len(page.dialogues)), cols=1)
            for row, dialog in zip(table.rows, page.dialogues):
                run = row.cells[0].paragraphs[0].add_run(f"{dialog.speaker_id}：{dialog.text}")
                run.font.name = "Microsoft JhengHei"
                run.font.size = Pt(16)
                run.bold = True
        doc.save(out)

    def _export_docx_word_shapes(self, episode: EpisodeManifest, out: Path) -> None:
        try:
            import pythoncom
            import win32com.client
        except ImportError as exc:
            raise ImportError("pywin32 unavailable") from exc
        pythoncom.CoInitialize()
        word = None
        doc = None
        try:
            word = win32com.client.DispatchEx("Word.Application")
            word.Visible = False
            word.DisplayAlerts = 0
            doc = word.Documents.Add()
            # 部分 zh-TW Word/pywin32 組合呼叫 CentimetersToPoints 會回傳
            # E_FAIL；直接使用 Word 的 point 單位可避免 locale-sensitive COM helper。
            doc.PageSetup.PageWidth = 595.276
            doc.PageSetup.PageHeight = 841.89
            anchors = [doc.Paragraphs(1).Range]
            selection = word.Selection
            for _ in episode.pages[1:]:
                selection.EndKey(6)
                selection.InsertBreak(7)
                selection.TypeParagraph()
                anchors.append(doc.Paragraphs(doc.Paragraphs.Count).Range)
            for index, (page, anchor) in enumerate(zip(episode.pages, anchors), start=1):
                background = doc.Shapes.AddPicture(
                    str(self.resolve_asset(episode, page.image_asset_id or "")),
                    False,
                    True,
                    18,
                    18,
                    559,
                    806,
                    anchor,
                )
                background.Name = f"P{index:02d}_BACKGROUND_FIXED"
                background.AlternativeText = "固定漫畫背景；請勿直接編輯文字於圖片內"
                background.RelativeHorizontalPosition = 1
                background.RelativeVerticalPosition = 1
                # AddPicture 會先以目前頁邊界解讀 Left/Top；切換成 relative-to-page
                # 後需重設座標，否則圖片仍會保留 Word 預設 margin 的位移。
                background.Left = 18
                background.Top = 18
                background.WrapFormat.Type = 3
                background.ZOrder(1)
                for bubble_no, dialog in enumerate(self.resolve_dialogue_layout(episode, page), start=1):
                    left = 18 + dialog.x * 559
                    top = 18 + dialog.y * 806
                    width = max(178, dialog.w * 559)
                    body_height = max(118, dialog.h * 806)
                    target_x = 18 + (dialog.tail_x or dialog.x + dialog.w / 2) * 559
                    target_y = 18 + (dialog.tail_y or dialog.y + dialog.h + 0.1) * 806
                    total_height = body_height
                    # Word msoShapeRoundedRectangularCallout：尾巴與圓角文字框是同一個 Shape。
                    shape = doc.Shapes.AddShape(106, left, top, width, total_height, anchor)
                    shape.Name = f"P{index:02d}_{bubble_no}_DIALOGUE_CALLOUT_EDITABLE"
                    shape.AlternativeText = "可直接編輯並移動的一體成形漫畫對話框"
                    shape.RelativeHorizontalPosition = 1
                    shape.RelativeVerticalPosition = 1
                    shape.Left = left
                    shape.Top = top
                    shape.WrapFormat.Type = 3
                    shape.Fill.ForeColor.RGB = 0xFFFFFF
                    shape.Line.ForeColor.RGB = (0x79 << 16) | (0x4E << 8) | 0x1F
                    shape.Line.Weight = 1.75
                    tip_ratio = (target_x - left) / max(1, width)
                    shape.Adjustments.SetItem(1, min(2.5, max(-2.5, tip_ratio - 0.5)))
                    # Adjustment 2 可大於 1，讓尾巴延伸到本體外，而不把文字框一起拉高。
                    shape.Adjustments.SetItem(2, min(5.5, max(0.62, (target_y - top) / max(1, body_height))))
                    shape.Adjustments.SetItem(3, 0.18)
                    shape.TextFrame.MarginLeft = 9
                    shape.TextFrame.MarginRight = 9
                    shape.TextFrame.MarginTop = 7
                    shape.TextFrame.MarginBottom = 7
                    shape.TextFrame.TextRange.Text = dialog.text
                    shape.TextFrame.TextRange.Font.Name = "Microsoft JhengHei"
                    shape.TextFrame.TextRange.Font.Size = min(18, max(13, dialog.font_size))
                    shape.TextFrame.TextRange.Font.Bold = True
                    shape.TextFrame.TextRange.Font.Color = 0x000000
                    shape.TextFrame.TextRange.ParagraphFormat.Alignment = 1
                    size = float(min(18, max(13, dialog.font_size)))
                    while bool(shape.TextFrame.Overflowing) and size > 12.0:
                        size -= 0.5
                        shape.TextFrame.TextRange.Font.Size = size
                    if bool(shape.TextFrame.Overflowing):
                        raise RuntimeError(f"Word bubble overflow: P{index:02d}_{bubble_no}")
            # Word COM 不接受相對路徑；即使 ComicStore 以相對 project root 初始化也要轉絕對路徑。
            doc.SaveAs2(str(out.resolve()), 16)
        except Exception as exc:
            raise RuntimeError(f"Word Shapes 匯出失敗: {exc}") from exc
        finally:
            if doc is not None:
                doc.Close(False)
            if word is not None:
                word.Quit()
            pythoncom.CoUninitialize()

    def create_source_zip(self, project_id: str, story_id: str, version: str) -> Path:
        episode = self.get_episode(project_id, story_id, version)
        root = self.episode_dir(project_id, story_id, version)
        self._sync_source_pack(episode)
        out = root / "exports" / f"{episode.story_id}_{version}_source_pack.zip"
        out.parent.mkdir(parents=True, exist_ok=True)
        with zipfile.ZipFile(out, "w", compression=zipfile.ZIP_DEFLATED) as zf:
            for path in sorted([root / "manifest.json", *(root / "source").glob("*.md")]):
                if path.is_file():
                    zf.write(path, path.relative_to(root).as_posix())
        return out

    def publish(self, project_id: str, story_id: str, version: str, *, published_by: str, channel: str = "internal_reader") -> EpisodeManifest:
        episode = self.get_episode(project_id, story_id, version)
        if episode.state != "CURRENT":
            raise ComicGateError("只有 CURRENT episode 可以發布")
        report = self.validate_episode(episode)
        if not report.publish_ready:
            raise ComicGateError("QA gate 未通過，不可發布")
        _, _ = self.export_html(project_id, story_id, version)
        current = self.get_episode(project_id, story_id, version)
        release = ReleaseRecord(
            release_id=f"release_{uuid.uuid4().hex[:8]}",
            channel=channel,
            public_version=version,
            published_by=published_by,
            url=f"/projects/{safe_id(project_id)}/comics/reader/{safe_id(story_id)}",
        )
        # CURRENT 平常不可 update；release 是不可變內容上的 audit append，直接受鎖保存。
        with self._lock:
            saved = current.model_copy(update={"releases": [*current.releases, release], "updated_at": utc_now(), "revision": current.revision + 1})
            return self._save_episode(saved)

    def withdraw_release(self, project_id: str, story_id: str, version: str, release_id: str) -> EpisodeManifest:
        episode = self.get_episode(project_id, story_id, version)
        found = False
        releases: list[ReleaseRecord] = []
        for release in episode.releases:
            if release.release_id == release_id and release.withdrawn_at is None:
                release = release.model_copy(update={"withdrawn_at": utc_now()})
                found = True
            releases.append(release)
        if not found:
            raise ComicNotFoundError(f"可撤回 release 不存在: {release_id}")
        with self._lock:
            saved = episode.model_copy(
                update={
                    "releases": releases,
                    "updated_at": utc_now(),
                    "revision": episode.revision + 1,
                }
            )
            return self._save_episode(saved)

    def _sync_source_pack(self, episode: EpisodeManifest) -> None:
        root = self.episode_dir(episode.project_id, episode.story_id, episode.version)
        source = root / "source"
        source.mkdir(parents=True, exist_ok=True)
        storyboard = [f"# {episode.title}｜Storyboard", ""]
        dialogue = [f"# {episode.title}｜Dialogue Script", ""]
        prompts = [f"# {episode.title}｜Image Prompts", ""]
        for page in episode.pages:
            storyboard.extend([
                f"## P{page.page_no:02d}｜{page.beat or '未命名節拍'}",
                f"- Scene: {page.scene_description}",
                f"- Camera: {page.camera}",
                f"- Learning point: {page.learning_point}",
                f"- Evidence: {', '.join(page.evidence_ids) or 'TBD'}",
                "",
            ])
            dialogue.append(f"## P{page.page_no:02d}")
            dialogue.extend(f"- {item.speaker_id}: {item.text}" for item in page.dialogues)
            dialogue.append("")
            prompts.extend([f"## P{page.page_no:02d}", "```text", page.image_prompt or "TBD", "```", ""])
        evidence = [f"# {episode.title}｜Technical Sources", ""]
        for item in episode.evidence:
            evidence.extend([
                f"## {item.source_id}｜{item.title}",
                f"- Type: {item.source_type}",
                f"- Publisher: {item.publisher}",
                f"- URL: {item.url}",
                f"- Citation: {item.citation}",
                f"- Supported claims: {'; '.join(item.supported_claims) or 'TBD'}",
                f"- Limits: {'; '.join(item.limits) or 'TBD'}",
                "",
            ])
        qa = [f"# {episode.title}｜QA Report", "", f"State: {episode.state}", ""]
        qa.extend(f"- {item.gate}: **{item.result}** — {item.evidence} ({item.reviewer}, {item.timestamp})" for item in episode.qa_records)
        revisions = [f"# {episode.title}｜Revision Notes", "", f"- revision: {episode.revision}", f"- updated_at: {episode.updated_at}", f"- state: {episode.state}"]
        files = {
            "storyboard.md": storyboard,
            "dialogue_script.md": dialogue,
            "image_prompts.md": prompts,
            "technical_sources.md": evidence,
            "qa_report.md": qa,
            "revision_notes.md": revisions,
        }
        for name, lines in files.items():
            (source / name).write_text("\n".join(lines).rstrip() + "\n", encoding="utf-8")


def decode_data_url(value: str) -> tuple[bytes, str]:
    match = re.fullmatch(r"data:(image/(?:png|jpeg|webp));base64,(.+)", value, re.DOTALL)
    if not match:
        raise ValueError("必須是 PNG/JPEG/WEBP base64 data URL")
    mime, payload = match.groups()
    try:
        data = base64.b64decode(payload, validate=True)
    except ValueError as exc:
        raise ValueError("base64 asset 無效") from exc
    suffix = {"image/png": ".png", "image/jpeg": ".jpg", "image/webp": ".webp"}[mime]
    return data, suffix


def make_mock_scene(page_no: int) -> str:
    """只供離線測試；圖上明確標 MOCK，不得通過 human approval。"""
    from PIL import Image, ImageDraw

    image = Image.new("RGB", (768, 1086), "#173042")
    draw = ImageDraw.Draw(image)
    draw.rectangle((48, 48, 720, 1038), outline="#49c8a5", width=8)
    draw.text((82, 90), f"MOCK SCENE {page_no:02d}", fill="#ffffff")
    draw.text((82, 140), "NOT A GENERATED OR REVIEWED COMIC ASSET", fill="#f5ae53")
    buf = io.BytesIO()
    image.save(buf, format="PNG")
    return "data:image/png;base64," + base64.b64encode(buf.getvalue()).decode("ascii")
